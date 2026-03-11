#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将 Customer360 同步稿 Markdown 转为 Word 文档"""

import re
import sys
import os

# 若在项目内安装了 python-docx
_venv = os.path.join(os.path.dirname(__file__), '..', 'venv_docx')
if os.path.isdir(_venv):
    sys.path.insert(0, _venv)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table(doc, rows, header=True):
    """添加表格到文档"""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text).strip()
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            if header and i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    return table

def parse_md(md_path):
    """简单解析 markdown，返回 (blocks) 列表，每个元素为 ('heading'|'para'|'list'|'table'|'hr', content)"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 水平线
        if stripped in ('---', '***', '___'):
            blocks.append(('hr', None))
            i += 1
            continue

        # 标题
        if stripped.startswith('#'):
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            title = stripped[level:].strip()
            blocks.append(('heading', (level, title)))
            i += 1
            continue

        # 表头分隔行（| --- | --- |）
        if stripped.startswith('|') and '---' in stripped:
            if blocks and blocks[-1][0] == 'para':
                prev = blocks.pop()
                header_line = prev[1]
            else:
                i += 1
                continue
            def split_table_row(s):
                parts = re.split(r'\|', s)
                if len(parts) < 2:
                    return [s.strip()]
                return [p.strip() for p in parts[1:-1]] if parts[0] == '' and parts[-1] == '' else [p.strip() for p in parts if p.strip()]
            header_cells = split_table_row(header_line)
            table_rows = [header_cells] if header_cells else []
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                cells = split_table_row(row_line)
                if cells:
                    table_rows.append(cells)
                i += 1
            if table_rows:
                blocks.append(('table', table_rows))
            continue

        # 表格行（无表头分隔的表格，这里不单独处理，交给上面）
        if stripped.startswith('|') and i > 0 and blocks and blocks[-1][0] == 'table':
            # 已在上方处理
            i += 1
            continue

        # 无序列表
        if stripped.startswith('- ') and not stripped.startswith('- ---'):
            items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item = lines[i].strip()[2:].strip()
                items.append(item)
                i += 1
            blocks.append(('list', items))
            continue

        # 有序列表
        if re.match(r'^\d+\.\s', stripped):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item = re.sub(r'^\d+\.\s', '', lines[i].strip())
                items.append(item)
                i += 1
            blocks.append(('list', items))
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落（可能后面紧跟表格）
        blocks.append(('para', stripped))
        i += 1

    return blocks

def md_to_docx(md_path, docx_path):
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    blocks = parse_md(md_path)

    for kind, content in blocks:
        if kind == 'hr':
            doc.add_paragraph('_' * 50)
            continue
        if kind == 'heading':
            level, title = content
            p = doc.add_heading(title, level=min(level, 3))
            continue
        if kind == 'para':
            doc.add_paragraph(content)
            continue
        if kind == 'list':
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
            continue
        if kind == 'table':
            add_table(doc, content, header=True)

    doc.save(docx_path)
    print('已生成:', docx_path)

if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base, 'Customer360_客户确认版_20260311_同步稿.md')
    docx_path = os.path.join(base, 'Customer360_客户确认版_20260311_同步稿.docx')
    md_to_docx(md_path, docx_path)
